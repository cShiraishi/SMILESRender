import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_table_border(table):
    """Adds a simple black border to the table if no styles are available."""
    tbl = table._tbl
    tblPr = tbl.xpath('w:tblPr')[0]
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def update_manuscript_docx(docx_path, json_data_path):
    print(f"Loading {docx_path}...")
    doc = Document(docx_path)
    
    doc.add_page_break()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Supplementary Material')
    run.bold = True
    run.font.size = Pt(16)
    
    p2 = doc.add_paragraph()
    run2 = p2.add_run('Table S1: Full List of RDKit Descriptors in SMILESRender')
    run2.bold = True
    run2.font.size = Pt(14)
    
    with open(json_data_path, 'r', encoding='utf-8') as f:
        descriptors = json.load(f)
    
    table = doc.add_table(rows=1, cols=4)
    try:
        table.style = 'Table Grid'
    except:
        print("Style 'Table Grid' not found, applying manual borders.")
        set_table_border(table)
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Description'
    hdr_cells[3].text = 'RDKit Function'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    
    for i, d in enumerate(descriptors, 1):
        if i % 10 == 0: print(f"Adding row {i}...")
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = str(d.get('Nome', ''))
        row_cells[2].text = str(d.get('Definição', ''))
        row_cells[3].text = str(d.get('Função RDKit', ''))
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    doc.save(docx_path)
    print(f"Successfully updated {docx_path} with {len(descriptors)} descriptors.")

if __name__ == "__main__":
    update_manuscript_docx(
        'SMILESRender_manuscript_final.docx', 
        'tmp/descriptors_data.json'
    )
