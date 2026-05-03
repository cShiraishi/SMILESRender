#!/usr/bin/env python3
"""
generate_manuscript_v5.py
Generates SMILESRender_JCheminform_2026_v5.docx from the markdown manuscript.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour constants ─────────────────────────────────────────────────────────
NAVY       = RGBColor(0x0D, 0x1F, 0x3C)
TEAL       = RGBColor(0x00, 0x7A, 0x6E)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
SHADED_BG  = RGBColor(0xEE, 0xF2, 0xF7)
ROW_ALT    = RGBColor(0xF8, 0xFA, 0xFC)
ROW_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helper: set paragraph shading ────────────────────────────────────────────
def set_paragraph_shading(para, fill_color: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(fill_color.red, fill_color.green, fill_color.blue)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

# ── Helper: set cell shading ─────────────────────────────────────────────────
def set_cell_shading(cell, fill_color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(fill_color.red, fill_color.green, fill_color.blue)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# ── Helper: set cell borders ─────────────────────────────────────────────────
def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '0D1F3C')
        tblBorders.append(border)
    tblPr.append(tblBorders)

# ── Helper: apply paragraph spacing ──────────────────────────────────────────
def set_para_spacing(para, space_before=0, space_after=4, line_spacing=None):
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if line_spacing:
        pf.line_spacing = line_spacing

# ── Helper: add run with font ────────────────────────────────────────────────
def add_run(para, text, bold=False, italic=False, size=10.5, color=None,
            font_name='Calibri', underline=False):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.name = font_name
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

# ── Heading helpers ───────────────────────────────────────────────────────────
def add_h1(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, space_before=12, space_after=6)
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    return para

def add_h2(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, space_before=10, space_after=4)
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    return para

def add_h3(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, space_before=8, space_after=4)
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = TEAL
    return para

def add_h4(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, space_before=6, space_after=3)
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = TEAL
    return para

def add_body(doc, text, space_after=4):
    para = doc.add_paragraph()
    set_para_spacing(para, space_after=space_after)
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    return para

def add_body_mixed(doc, parts, space_after=4):
    """parts: list of (text, bold, italic)"""
    para = doc.add_paragraph()
    set_para_spacing(para, space_after=space_after)
    for text, bold, italic in parts:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
    return para

def add_italic_body(doc, text, space_after=4):
    para = doc.add_paragraph()
    set_para_spacing(para, space_after=space_after)
    run = para.add_run(text)
    run.italic = True
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    return para

# ── Table caption ─────────────────────────────────────────────────────────────
def add_table_caption(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, space_before=3, space_after=8)
    run = para.add_run(text)
    run.italic = True
    run.font.name = 'Calibri'
    run.font.size = Pt(9)

# ── Styled table builder ──────────────────────────────────────────────────────
def build_styled_table(doc, headers, rows, col_widths=None):
    """Build a Table Grid table with NAVY header and alternating rows."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths if provided
    if col_widths:
        for i, col in enumerate(table.columns):
            col.width = Cm(col_widths[i])

    # Header row
    hdr_row = table.rows[0]
    for j, hdr in enumerate(headers):
        cell = hdr_row.cells[j]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = para.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after  = Pt(2)
        run = para.add_run(hdr)
        run.bold       = True
        run.font.name  = 'Calibri'
        run.font.size  = Pt(9)
        run.font.color.rgb = WHITE

    # Data rows
    for i, row_data in enumerate(rows):
        tr = table.rows[i + 1]
        bg = ROW_ALT if (i % 2 == 0) else ROW_WHITE
        for j, cell_text in enumerate(row_data):
            cell = tr.cells[j]
            set_cell_shading(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            pf = para.paragraph_format
            pf.space_before = Pt(2)
            pf.space_after  = Pt(2)
            # Centre-align if it looks like a checkmark/symbol column
            if str(cell_text).strip() in ['✓', '—', '✓ᵃ', 'Partial', 'Partialᵇ',
                                           'Partialᶜ', 'Partialᵈ', 'Partialᵉ', 'Partialᶠ',
                                           'Partialᵍ', '—ᵉ', 'Manualᵍ']:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Parse inline bold (text wrapped in **...**)
            cell_str = str(cell_text)
            _add_mixed_run_to_para(para, cell_str, size=9)

    set_table_borders(table)
    return table

def _add_mixed_run_to_para(para, text, size=9, default_bold=False):
    """Add runs to para, handling **bold** markers."""
    import re
    segments = re.split(r'(\*\*[^*]+\*\*)', text)
    for seg in segments:
        if seg.startswith('**') and seg.endswith('**'):
            run = para.add_run(seg[2:-2])
            run.bold = True
        else:
            run = para.add_run(seg)
            run.bold = default_bold
        run.font.name = 'Calibri'
        run.font.size = Pt(size)


# ── Figure legend box ─────────────────────────────────────────────────────────
def add_figure_legend(doc, fig_num, legend_text):
    """Shaded single-cell table + bold caption below."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, SHADED_BG)
    para = cell.paragraphs[0]
    pf = para.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after  = Pt(4)
    run = para.add_run(f"[Figure {fig_num} — placeholder for artwork]")
    run.italic = True
    run.font.name = 'Calibri'
    run.font.size = Pt(10)

    # Bold caption below
    cap = doc.add_paragraph()
    set_para_spacing(cap, space_before=3, space_after=10)
    r = cap.add_run(f"Figure {fig_num}. ")
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r2 = cap.add_run(legend_text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10)


# ── ESOL blockquote ───────────────────────────────────────────────────────────
def add_esol_block(doc):
    """Indented shaded blockquote, Courier New 10pt."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, SHADED_BG)
    para = cell.paragraphs[0]
    pf = para.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after  = Pt(4)
    pf.left_indent  = Inches(0.2)
    run = para.add_run("log S = 0.16 − 0.63·cLogP − 0.0062·MW + 0.066·RotB − 0.74·AP")
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    # small space after block
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


# ── Reference helper ──────────────────────────────────────────────────────────
def add_reference(doc, num, text_parts):
    """text_parts: list of (text, bold, italic)"""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after  = Pt(3)
    pf.left_indent  = Inches(0.3)
    pf.first_line_indent = Inches(-0.3)
    for text, bold, italic in text_parts:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Calibri'
        run.font.size = Pt(9.5)


# ════════════════════════════════════════════════════════════════════════════════
# MAIN BUILD
# ════════════════════════════════════════════════════════════════════════════════
def build_document():
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

    # ── Default style ─────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # ════════════════════════════════════════════════════════════════════════
    # TITLE BLOCK
    # ════════════════════════════════════════════════════════════════════════
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(title_para, space_before=0, space_after=10)
    r = title_para.add_run(
        "SMILESRender: A Unified Open-Source Web Platform for Centralized Cheminformatics, "
        "Multi-Engine ADMET Profiling, and Molecular Analysis"
    )
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(14)
    r.font.color.rgb = NAVY

    # Authors
    auth = doc.add_paragraph()
    auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(auth, space_after=4)
    add_run(auth, "Authors: ", bold=True, size=10.5)
    add_run(auth, "Rui A. B. Shiraishi¹*, Gabriel Grechuk¹", size=10.5)

    # Affiliations
    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(aff, space_after=4)
    add_run(aff, "Affiliations: ", bold=True, size=10.5)
    add_run(aff, "¹ [Department], [Institution], [City, Country]", size=10.5)

    # Corresponding
    corr = doc.add_paragraph()
    corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(corr, space_after=4)
    add_run(corr, "*Corresponding author: ", bold=True, size=10.5)
    add_run(corr, "carlos.seiti.shiraishi@gmail.com", size=10.5)

    # Journal / Date / Keywords
    jrnl = doc.add_paragraph()
    jrnl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(jrnl, space_after=4)
    add_run(jrnl, "Target journal: ", bold=True, size=10.5)
    add_run(jrnl, "Journal of Cheminformatics — Software Article  |  Submitted: May 2026", size=10.5)

    kw = doc.add_paragraph()
    kw.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(kw, space_after=12)
    add_run(kw, "Keywords: ", bold=True, size=10.5)
    add_run(kw, "cheminformatics; drug discovery; ADMET; open-source; web platform; molecular visualization; "
               "workflow integration; blood-brain barrier; toxicity prediction; reproducibility; DataWarrior; KNIME", size=10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ════════════════════════════════════════════════════════════════════════
    # ABSTRACT
    # ════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Abstract")

    add_body_mixed(doc, [("Background: ", True, False),
        ("Modern drug discovery demands proficiency across a fragmented set of computational tools: "
         "molecular visualization software, ADMET prediction servers, descriptor calculators, structure editors, "
         "and similarity search engines. Researchers navigate this landscape by maintaining parallel installations "
         "of desktop applications (DataWarrior, MarvinSketch), programming environments (KNIME, RDKit scripts), "
         "and browser sessions across multiple disconnected web services — a workflow that imposes substantial "
         "context-switching overhead and fundamentally limits reproducibility. No single open-source platform "
         "currently consolidates the full cheminformatics stack — structure rendering, interactive drawing, "
         "multi-engine ADMET profiling, structural alert screening, QSAR-ready descriptor computation, chemical "
         "similarity search, and automated risk interpretation — within one session-consistent, containerised, "
         "offline-capable deployment.", False, False)])

    add_body_mixed(doc, [("Results: ", True, False),
        ("We present SMILESRender, an open-source web-based cheminformatics hub built on a hybrid architecture. "
         "The platform delivers: (i) high-quality 2D molecular rendering and reaction visualization; "
         "(ii) an embedded JSME structure editor; (iii) comprehensive ADMET profiling combining three local machine "
         "learning models (Tox21 multi-endpoint, BBB permeability, and 53-property Chemprop D-MPNN via ", False, False),
        ("admet_ai", False, True),
        (") with three external oracle services (StopTox, StopLight, ProTox 3.0) — covering over 85 ADMET endpoints; "
         "(iv) structural alert screening (PAINS, BRENK, NIH catalogs); (v) 60+ local RDKit descriptors with four "
         "fingerprint export formats; (vi) chemical similarity search and IUPAC nomenclature conversion; and "
         "(vii) a rule-based Automated Interpretation Engine converting numerical outputs into severity-classified "
         "plain-language narratives. A systematic benchmark against five major open-source cheminformatics platforms "
         "— DataWarrior, KNIME, Galaxy cheminformatics, ChemMine Tools, and MarvinSketch — demonstrates that "
         "SMILESRender is the only open-source tool providing web-native, integrated, multi-engine ADMET profiling "
         "with local ML models and automated interpretation. Batch processing of 20 thieno[2,3-b]pyridine derivatives "
         "completed in under 12 minutes versus 2 h 52 min for the equivalent manual multi-tool workflow.", False, False)])

    add_body_mixed(doc, [("Conclusions: ", True, False),
        ("SMILESRender addresses the operational fragmentation that limits reproducibility and accessibility in "
         "computational medicinal chemistry, consolidating into a single Docker-deployable platform capabilities "
         "that previously required five or more separate tools. Source code is available under the MIT license at "
         "https://github.com/rubithedev/smiles-render-web; a public cloud instance runs at "
         "https://smiles-render.onrender.com.", False, False)])

    # ════════════════════════════════════════════════════════════════════════
    # BACKGROUND
    # ════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Background")
    add_h2(doc, "The Fragmentation Problem in Computational Drug Discovery")

    add_body(doc,
        "The computational toolkit of a medicinal chemist typically spans multiple disconnected environments. "
        "A typical ADMET evaluation workflow for a 20-compound library involves: drawing or editing structures "
        "in a dedicated editor (MarvinSketch, ChemDoodle, or JSME), rendering 2D images for presentation "
        "(DataWarrior, ChemDraw, or KNIME), computing physicochemical descriptors (RDKit scripts, DataWarrior, "
        "or SwissADME), querying acute toxicity predictions (StopTox or ProTox), querying ADMET predictions "
        "(pkCSM, ADMETlab, or admet_ai), checking structural alerts (custom scripts or SwissADME), and finally "
        "consolidating results across incompatible spreadsheets.")

    add_body(doc,
        "This fragmentation is not merely inconvenient — it is a reproducibility risk. Different tools implement "
        "the same descriptors with different normalization conventions. Batch sizes differ between services. Results "
        "are downloaded in incompatible formats. Session state is lost when browsers are closed. And any external "
        "service can change its API, enforce rate limits, or go offline between the time a paper is submitted and "
        "when a reviewer attempts to reproduce the analysis.")

    add_body(doc,
        "The problem is particularly acute for synthetic chemists and biologists — the majority of medicinal "
        "chemistry researchers — who are not computational specialists. For these users, the overhead of setting "
        "up and maintaining parallel tools can make computational ADMET screening practically inaccessible, "
        "despite the availability of excellent free prediction services.")

    add_h2(doc, "Existing Open-Source Tools and Their Limitations")

    add_body(doc,
        "The cheminformatics community has produced a rich set of open-source tools, each addressing a specific "
        "aspect of the workflow:")

    # DataWarrior
    p = doc.add_paragraph()
    set_para_spacing(p, space_after=4)
    add_run(p, "DataWarrior", bold=True, size=10.5)
    add_run(p, " [1] (Actelion/Sanofi, freely distributed) is the most comprehensive open-source desktop "
               "cheminformatics application, providing 2D/3D visualization, compound clustering, self-organising "
               "map (SOM) analysis, activity cliff detection, and basic physicochemical property calculation. "
               "However, DataWarrior is desktop-only (Windows/macOS/Linux installation required), provides no "
               "web interface, cannot query external ADMET prediction services, and its built-in ADMET coverage is "
               "limited to basic drug-likeness filters (Lipinski, Veber) and a small set of local predictions. "
               "It has no BBB permeability model, no multi-endpoint toxicity profiling, and no automated "
               "interpretation layer.", size=10.5)

    # KNIME
    p = doc.add_paragraph()
    set_para_spacing(p, space_after=4)
    add_run(p, "KNIME", bold=True, size=10.5)
    add_run(p, " [2] with cheminformatics extensions (RDKit KNIME nodes, CDK nodes) is an extremely powerful "
               "visual programming platform capable of building sophisticated cheminformatics workflows. Its "
               "strength is flexibility: experts can connect any combination of nodes to build custom pipelines. "
               "Its weakness is accessibility: KNIME requires significant programming knowledge to configure "
               "effectively, has a steep learning curve, and provides no built-in ADMET prediction — external "
               "API nodes must be added and maintained separately. It is not web-native and is not suitable for "
               "non-computational users.", size=10.5)

    # Galaxy
    p = doc.add_paragraph()
    set_para_spacing(p, space_after=4)
    add_run(p, "Galaxy", bold=True, size=10.5)
    add_run(p, " [3] (cheminformatics tools at usegalaxy.eu) provides server-based reproducible scientific "
               "workflows, including some cheminformatics tools (structure conversion, SMILES processing). Galaxy "
               "excels at reproducible data analysis pipelines but has very limited cheminformatics capabilities, "
               "no dedicated ADMET prediction tools, and requires workflow creation expertise.", size=10.5)

    # ChemMine
    p = doc.add_paragraph()
    set_para_spacing(p, space_after=4)
    add_run(p, "ChemMine Tools", bold=True, size=10.5)
    add_run(p, " [4] is a web-based cheminformatics platform offering compound comparison, structural clustering, "
               "and molecular similarity search. It provides no ADMET prediction, no descriptor export, and no "
               "structural alert screening.", size=10.5)

    # MarvinSketch
    p = doc.add_paragraph()
    set_para_spacing(p, space_after=4)
    add_run(p, "MarvinSketch", bold=True, size=10.5)
    add_run(p, " (ChemAxon, free academic licence) provides excellent structure drawing and a subset of "
               "physicochemical property calculations (pKa, logP, solubility). However, it is not fully "
               "open-source, provides no ADMET prediction beyond basic physicochemical parameters, and offers "
               "no batch processing pipeline.", size=10.5)

    # RDKit
    p = doc.add_paragraph()
    set_para_spacing(p, space_after=4)
    add_run(p, "RDKit", bold=True, size=10.5)
    add_run(p, " [5] is the community standard for programmatic cheminformatics — it provides exactly the "
               "capabilities needed for descriptor computation, fingerprint generation, and structural alert "
               "screening. However, RDKit is a Python/C++ library, not a user-facing application. Using it "
               "requires programming expertise, and it provides no ADMET prediction models of its own.", size=10.5)

    add_body(doc,
        "A critical gap therefore exists: there is no open-source, web-native platform that consolidates "
        "structure rendering, interactive editing, multi-endpoint ADMET profiling, structural alert screening, "
        "descriptor computation, and automated interpretation in a single session-consistent deployment accessible "
        "without programming expertise.")

    add_h2(doc, "Centralisation as a Scientific and Practical Imperative")

    add_body(doc,
        "Beyond convenience, centralisation of cheminformatics workflows has direct scientific implications:")

    p = doc.add_paragraph()
    set_para_spacing(p, space_after=4)
    add_run(p, "Reproducibility: ", bold=True, size=10.5)
    add_run(p, "When all computations for a study are performed within a single containerised deployment, the "
               "full computational environment can be captured and shared. This is qualitatively different from "
               "a methods section that lists five separate tools, each with its own version dependencies and "
               "update history.", size=10.5)

    p = doc.add_paragraph()
    set_para_spacing(p, space_after=4)
    add_run(p, "Session consistency: ", bold=True, size=10.5)
    add_run(p, "When SMILES strings are entered once and propagated to all tools within the same session, there "
               "is no risk of transcription errors, copy-paste mistakes, or tautomer/canonicalisation differences "
               "between tools.", size=10.5)

    p = doc.add_paragraph()
    set_para_spacing(p, space_after=4)
    add_run(p, "Accessibility: ", bold=True, size=10.5)
    add_run(p, "Web-native interfaces eliminate installation barriers, enabling computational ADMET access for "
               "synthetic chemists, pharmacologists, and students without programming backgrounds.", size=10.5)

    p = doc.add_paragraph()
    set_para_spacing(p, space_after=4)
    add_run(p, "Integrated interpretation: ", bold=True, size=10.5)
    add_run(p, "When predictions from multiple tools are available simultaneously, cross-tool consensus (e.g., "
               "BBB classification concordant between local GBM model and Chemprop BBB_Martins) provides more "
               "reliable signals than any single-tool prediction in isolation.", size=10.5)

    add_body(doc,
        "SMILESRender was built to address this gap within the constraints of open-source, MIT-licensed software "
        "deployable on any infrastructure.")

    # ════════════════════════════════════════════════════════════════════════
    # IMPLEMENTATION
    # ════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Implementation")
    add_h2(doc, "System Architecture")

    add_body_mixed(doc, [
        ("SMILESRender follows a three-tier hybrid architecture (Figure 1). A React 19/TypeScript single-page "
         "application communicates with a Python Flask 3.0 backend (Waitress 3.0 WSGI server). The backend "
         "separates two computation pathways: ", False, False),
        ("(i) local in-process computation", False, True),
        (" via RDKit 2024.3.6, scikit-learn 1.8, and ", False, False),
        ("admet_ai", False, True),
        (", requiring no network access; and ", False, False),
        ("(ii) external oracle orchestration", False, True),
        (" via asynchronous proxy requests to StopTox, StopLight, and ProTox 3.0, each isolated in a "
         "ToolErrorBoundary.", False, False)])

    add_body(doc,
        "Because all three local ML models run independently of the external orchestration layer, a minimum viable "
        "ADMET profile is always available regardless of network conditions. A Redis 7.4 cache (24-hour TTL, keyed "
        "by MD5 of canonical SMILES) reduces redundant external calls by 60–80% in iterative workflows. Docker "
        "Compose containerises all three services (web server, Redis, Celery worker), ensuring bit-identical results "
        "across deployments. The backend exposes 19 REST endpoints across four namespaces (/render/*, /predict/*, "
        "/descriptors, /convert/*).")

    add_h2(doc, "Module 1 — Molecular Structure Rendering")

    add_body(doc,
        "SMILES strings are converted to 2D structural images via RDKit Draw.MolToImage with rdDepictor coordinate "
        "generation. Transparent-background PNG images are produced by alpha-channel replacement. Batch mode accepts "
        "up to 20 SMILES per request and returns a deduplicated ZIP archive. Supported export formats: PNG, JPEG, "
        "WEBP, TIFF, BMP, GIF, EPS, ICO. Reaction SMILES (reactants>>products) are handled via "
        "rdkit.Chem.Draw.ReactionToImage with full atom-mapping support.")

    add_body_mixed(doc, [
        ("Interactive structure drawing is provided via the JSME Molecular Editor [6], embedded as a "
         "browser-native JavaScript component with no installation requirement. JSME exports canonical SMILES "
         "that feed directly into the prediction pipeline.", False, False)])

    add_h2(doc, "Module 2 — Comprehensive ADMET Profiling")

    add_body(doc,
        "The ADMET module is the scientific centrepiece of SMILESRender, covering over 85 endpoints across all "
        "five ADMET categories through three complementary prediction layers.")

    add_h3(doc, "2.1 Local Machine Learning Models (100% available offline)")

    add_body_mixed(doc, [
        ("Tox21 Multi-Endpoint Toxicity (12 endpoints):", True, False),
        (" A Multi-Output Random Forest (ECFP4, 1,024 bits; scikit-learn 1.8) covers all 12 Tox21 Challenge "
         "bioassay endpoints, providing in vitro surrogates for nuclear receptor disruption and stress pathway "
         "activation. Mean AUC-ROC = 0.81 (5-fold stratified CV), consistent with published RF baselines.", False, False)])

    add_body_mixed(doc, [
        ("Blood-Brain Barrier Permeability:", True, False),
        (" A GradientBoosting classifier (ECFP4 2,048 bits + 9 pharmacokinetic descriptors; trained on curated "
         "B3DB, n = 7,643) predicts BBB+ or BBB− status with AUC-ROC = 0.92 on stratified hold-out (95% bootstrap "
         "CI [0.90, 0.94]). Each prediction is accompanied by a Tanimoto applicability domain (AD) flag "
         "(nearest-neighbour threshold 0.30) alerting when the compound is outside the training chemical space.", False, False)])

    add_body_mixed(doc, [
        ("Deep ADMET — 53 Properties (Chemprop D-MPNN via ", True, False),
        ("admet_ai", True, True),
        ("):", True, False),
        (" Pre-trained Chemprop Directed Message Passing Neural Network [7,8] models covering absorption, "
         "distribution, metabolism, excretion, and toxicity (full coverage in Section 2.2). Median AUC-ROC = 0.894 "
         "across 28 classification tasks (TDC leaderboard, Swanson et al. [8]). Mean inference: 280 ms per compound "
         "on a 4-core CPU.", False, False)])

    add_h3(doc, "2.2 External Oracle Services (supplementary, fault-isolated)")

    bullets_external = [
        ("StopTox", " [9]: six acute systemic toxicity endpoints (oral/dermal/inhalation LD50, eye irritation, "
         "skin sensitisation, aquatic toxicity); validated QSAR models from NIH/NTP."),
        ("StopLight", " [10]: eleven multi-parameter optimisation (MPO) scores for lead optimisation."),
        ("ProTox 3.0", " [11]: twelve organ-toxicity predictions (see ADMET section for full endpoint list)."),
    ]
    for bold_part, rest in bullets_external:
        p = doc.add_paragraph(style='List Bullet')
        set_para_spacing(p, space_after=3)
        add_run(p, bold_part, bold=True, size=10.5)
        add_run(p, rest, size=10.5)

    add_h3(doc, "2.3 ADMET Endpoint Coverage by Category")

    add_body(doc,
        "The following subsections describe the biological and clinical rationale for each ADMET endpoint covered "
        "by SMILESRender, grouped by category.")

    # Absorption
    p = doc.add_paragraph()
    set_para_spacing(p, space_after=4)
    add_run(p, "Absorption ", bold=True, size=10.5)
    add_run(p, "determines whether an orally administered compound reaches systemic circulation in adequate "
               "concentrations. Key endpoints covered:", size=10.5)

    absorption_bullets = [
        ("Human intestinal absorption (HIA):", " fraction absorbed via passive diffusion and active transport "
         "through the intestinal epithelium. HIA < 30% indicates poor oral bioavailability."),
        ("Caco-2 permeability:", " permeability through Caco-2 cell monolayers, a standard surrogate for "
         "intestinal permeability. Low Caco-2 (< 10⁻⁶ cm/s) correlates with poor HIA."),
        ("PAMPA permeability:", " passive transcellular permeability. Complements Caco-2 by isolating passive "
         "from active transport contributions."),
        ("P-glycoprotein (P-gp) substrate and inhibitor:", " P-gp is a major efflux transporter expressed at "
         "the intestinal epithelium, BBB, liver, and kidney. P-gp substrates face active efflux that reduces "
         "bioavailability and CNS penetration."),
        ("Oral bioavailability (F20%, F30%):", " fraction of the administered oral dose reaching systemic "
         "circulation. Values below the threshold indicate formulation challenges."),
        ("TPSA:", " computed locally via RDKit. TPSA > 140 Å² predicts poor passive intestinal absorption "
         "(Veber et al. [12])."),
    ]
    for bold_part, rest in absorption_bullets:
        p = doc.add_paragraph(style='List Bullet')
        set_para_spacing(p, space_after=2)
        add_run(p, bold_part, italic=True, size=10.5)
        add_run(p, rest, size=10.5)

    # Distribution
    p = doc.add_paragraph()
    set_para_spacing(p, space_before=6, space_after=4)
    add_run(p, "Distribution ", bold=True, size=10.5)
    add_run(p, "governs how a compound partitions between blood, tissues, and target organs following absorption:", size=10.5)

    distribution_bullets = [
        ("Blood-brain barrier (BBB) permeability:", " the single most pharmacologically critical distribution "
         "endpoint for CNS drug candidates. SMILESRender provides two independent BBB predictions — the local GBM "
         "model (B3DB-trained) and Chemprop BBB_Martins — enabling cross-model consensus assessment. Concordant "
         "BBB− from both models constitutes a high-confidence CNS-impermeability flag."),
        ("Plasma protein binding (PPBR):", " fraction bound to plasma proteins (albumin, α1-acid glycoprotein). "
         "High PPBR (> 95%) limits free drug concentration available for target binding."),
        ("Volume of distribution at steady state (VDss):", " reflects tissue partitioning. Low VDss indicates "
         "plasma-confined distribution; high VDss indicates extensive tissue uptake."),
        ("P-gp substrate:", " relevant at the BBB and other tissue barriers beyond intestinal absorption."),
    ]
    for bold_part, rest in distribution_bullets:
        p = doc.add_paragraph(style='List Bullet')
        set_para_spacing(p, space_after=2)
        add_run(p, bold_part, italic=True, size=10.5)
        add_run(p, rest, size=10.5)

    # Metabolism
    p = doc.add_paragraph()
    set_para_spacing(p, space_before=6, space_after=4)
    add_run(p, "Metabolism ", bold=True, size=10.5)
    add_run(p, "— primarily hepatic CYP450-mediated biotransformation — determines the rate of drug elimination "
               "and potential for drug–drug interactions (DDIs):", size=10.5)

    metabolism_bullets = [
        ("CYP1A2 inhibition/substrate:", " major route for caffeine, theophylline, clozapine. Inhibition causes "
         "DDIs with narrow-therapeutic-index substrates."),
        ("CYP2C9 inhibition/substrate:", " critical for warfarin, phenytoin, NSAIDs. CYP2C9 inhibition carries "
         "significant bleeding risk."),
        ("CYP2C19 inhibition/substrate:", " involved in clopidogrel activation (prodrug). Inhibition impairs "
         "antiplatelet efficacy."),
        ("CYP2D6 inhibition/substrate:", " metabolises 25% of marketed drugs including codeine, tamoxifen, "
         "antidepressants. Genetic polymorphisms create poor/ultra-rapid metaboliser populations."),
        ("CYP3A4 inhibition/substrate:", " the most important CYP isoform, responsible for ~50% of drug "
         "metabolism. Inhibition raises plasma levels of co-administered substrates."),
        ("Metabolic half-life (T1/2):", " determines dosing frequency. Short T1/2 (< 1 h) requires frequent "
         "dosing; very long T1/2 risks accumulation."),
        ("CYP polypharmacology flag (local interpretation engine):", " when ≥ 3 of 5 isoforms show inhibition "
         "probability ≥ 0.50, SMILESRender flags high DDI liability."),
    ]
    for bold_part, rest in metabolism_bullets:
        p = doc.add_paragraph(style='List Bullet')
        set_para_spacing(p, space_after=2)
        add_run(p, bold_part, italic=True, size=10.5)
        add_run(p, rest, size=10.5)

    # Excretion
    p = doc.add_paragraph()
    set_para_spacing(p, space_before=6, space_after=4)
    add_run(p, "Excretion ", bold=True, size=10.5)
    add_run(p, "governs drug clearance from the body:", size=10.5)

    excretion_bullets = [
        ("Hepatocyte clearance:", " intrinsic hepatic metabolic clearance. High clearance indicates rapid liver "
         "elimination and a need for frequent dosing or prodrug strategies."),
        ("Microsome clearance:", " microsomal metabolic clearance, a faster in vitro assay for oxidative metabolism."),
    ]
    for bold_part, rest in excretion_bullets:
        p = doc.add_paragraph(style='List Bullet')
        set_para_spacing(p, space_after=2)
        add_run(p, bold_part, italic=True, size=10.5)
        add_run(p, rest, size=10.5)

    # Toxicity
    p = doc.add_paragraph()
    set_para_spacing(p, space_before=6, space_after=4)
    add_run(p, "Toxicity ", bold=True, size=10.5)
    add_run(p, "— the most complex ADMET category — spans multiple mechanisms from direct organ damage to "
               "genotoxicity and regulatory risk:", size=10.5)

    toxicity_bullets = [
        ("hERG cardiotoxicity:", " inhibition of the hERG cardiac potassium channel causes QT interval "
         "prolongation, potentially leading to life-threatening arrhythmias (torsades de pointes). The hERG "
         "liability was responsible for withdrawal of multiple marketed drugs (terfenadine, cisapride, "
         "grepafloxacin). SMILESRender flags predicted hERG inhibition probability ≥ 0.40 as high risk, "
         "consistent with ICH E14 guidance."),
        ("Drug-induced liver injury (DILI):", " the leading cause of post-approval drug withdrawal. DILI "
         "prediction from SMILES is inherently difficult due to multi-mechanistic aetiology; the Chemprop model "
         "provides a probabilistic estimate that is appropriate for early screening."),
        ("AMES mutagenicity:", " in vitro bacterial reverse-mutation assay surrogate. Positive AMES is a "
         "regulatory concern under ICH S2(R1) and is required for NCE regulatory packages."),
        ("Carcinogenicity:", " long-term in vivo carcinogenicity risk. Valuable for early de-prioritisation of leads."),
        ("ClinTox:", " binary clinical toxicity flag derived from FDA-approved drug vs. clinical trial failure data."),
        ("LD50 (acute oral toxicity):", " estimated lethal dose in rodents. Classified by StopTox according to GHS: "
         "critical (< 50 mg/kg), high (50–300 mg/kg), moderate (300–2,000 mg/kg), low (> 2,000 mg/kg)."),
        ("Organ-specific toxicity (ProTox 3.0):", " neurotoxicity, nephrotoxicity, cardiotoxicity, "
         "immunotoxicity, cytotoxicity, hepatotoxicity, respiratory, and ecotoxicity — 12 endpoints covering "
         "target organs assessed in regulatory toxicology studies."),
        ("Tox21 12-endpoint in vitro panel:", " nuclear receptor activity (NR-AR, NR-AR-LBD, NR-AhR, "
         "NR-Aromatase, NR-ER, NR-ER-LBD, NR-PPAR-gamma) and stress response pathways (SR-ARE, SR-ATAD5, "
         "SR-HSE, SR-MMP, SR-p53). These endpoints are directly relevant to endocrine disruption screening and "
         "are increasingly required in regulatory submissions under REACH and EPA guidelines."),
    ]
    for bold_part, rest in toxicity_bullets:
        p = doc.add_paragraph(style='List Bullet')
        set_para_spacing(p, space_after=2)
        add_run(p, bold_part, italic=True, size=10.5)
        add_run(p, rest, size=10.5)

    # ── TABLE 1 ───────────────────────────────────────────────────────────────
    add_body(doc, "")  # spacer
    p_cap_title = doc.add_paragraph()
    set_para_spacing(p_cap_title, space_before=8, space_after=4)
    r = p_cap_title.add_run("Table 1. Complete ADMET endpoint coverage in SMILESRender by category and source.")
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(10)

    t1_headers = ["Category", "Endpoint", "Source", "Local/External"]
    t1_rows = [
        ["Absorption", "HIA, Caco-2, PAMPA, P-gp substrate/inhibitor, F20%, F30%", "admet_ai", "Local"],
        ["Absorption", "TPSA-based flag, Lipinski/Veber/Ghose/Egan/Muegge", "RDKit", "Local"],
        ["Distribution", "BBB (GBM model) + AD flag", "Local GBM", "Local"],
        ["Distribution", "BBB_Martins, PPBR, VDss", "admet_ai", "Local"],
        ["Metabolism", "CYP1A2/2C9/2C19/2D6/3A4 inhibition+substrate, T1/2", "admet_ai", "Local"],
        ["Excretion", "Hepatocyte clearance, microsome clearance", "admet_ai", "Local"],
        ["Toxicity", "hERG, DILI, AMES, carcinogenicity, ClinTox, LD50", "admet_ai", "Local"],
        ["Toxicity", "NR-AR, NR-AR-LBD, NR-AhR, NR-Aromatase, NR-ER, NR-ER-LBD, NR-PPAR-gamma, "
                     "SR-ARE, SR-ATAD5, SR-HSE, SR-MMP, SR-p53", "Tox21-RF", "Local"],
        ["Toxicity", "Oral/dermal/inhalation LD50, eye irritation, skin sensitisation, aquatic toxicity",
         "StopTox", "External"],
        ["Toxicity", "DILI, neurotoxicity, nephrotoxicity, cardiotoxicity, carcinogenicity, mutagenicity, "
                     "immunotoxicity, cytotoxicity, BBB, respiratory, ecotoxicity, clinical toxicity",
         "ProTox 3.0", "External"],
        ["Optimisation", "11 MPO scores (lead-likeness, CNS-MPO, PO score)", "StopLight", "External"],
        ["Structural alerts", "PAINS (A/B/C), BRENK, NIH catalogs", "RDKit FilterCatalog", "Local"],
    ]
    build_styled_table(doc, t1_headers, t1_rows, col_widths=[2.8, 7.0, 2.5, 2.5])
    add_table_caption(doc,
        "Total: 85+ distinct ADMET endpoints across all five categories. All local endpoints are available "
        "offline (100% uptime); external endpoints supplement with additional coverage when network is available.")

    # ── Flag-threshold table (Module 3) ───────────────────────────────────────
    add_h2(doc, "Module 3 — Automated Interpretation Engine")

    add_body_mixed(doc, [
        ("A rule-based engine (", False, False),
        ("admet_interpreter.py", False, True),
        (") converts the aggregated multi-tool numerical output into structured per-molecule risk profiles: "
         "severity-classified flags (low/moderate/high/critical), an overall risk level, and a plain-language "
         "narrative paragraph. This layer is particularly valuable for non-computational users who need actionable "
         "guidance rather than raw probability scores.", False, False)])

    add_body(doc,
        "Flag logic is grounded in established regulatory and pharmacological guidelines: GHS LD50 classification "
        "thresholds, ICH E14 hERG guidance, Veber TPSA absorption criteria [12], and Baell/Holloway PAINS "
        "definitions [13]. The engine explicitly communicates that ML probability outputs are relative discriminative "
        "scores, not calibrated absolute risk estimates, via a disclaimer embedded in every narrative.")

    # Flag-threshold table
    p_cap_title2 = doc.add_paragraph()
    set_para_spacing(p_cap_title2, space_before=8, space_after=4)
    r2 = p_cap_title2.add_run("Flag thresholds applied by the Automated Interpretation Engine.")
    r2.bold = True
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10)

    ft_headers = ["Parameter", "Low", "Moderate", "High", "Critical", "Guideline"]
    ft_rows = [
        ["hERG probability", "< 0.25", "0.25–0.40", "0.40–0.70", "≥ 0.70", "ICH E14"],
        ["DILI probability", "< 0.30", "0.30–0.50", "0.50–0.75", "≥ 0.75", "FDA DILI guidance"],
        ["LD50 (GHS class)", "> 2,000 mg/kg", "300–2,000", "50–300", "< 50 mg/kg", "GHS Rev.9"],
        ["TPSA", "< 90 Å²", "90–120 Å²", "120–140 Å²", "> 140 Å²", "Veber et al. [12]"],
        ["AMES", "—", "—", "Positive", "—", "ICH S2(R1)"],
        ["PAINS alerts", "0", "1", "2", "≥ 3", "Baell & Holloway [13]"],
        ["CYP inhibited (of 5)", "0", "1–2", "3–4", "5", "Internal heuristic"],
        ["Lipinski violations", "0", "1", "2", "≥ 3", "Lipinski et al. [15]"],
    ]
    build_styled_table(doc, ft_headers, ft_rows, col_widths=[3.0, 2.2, 2.2, 2.2, 2.2, 3.0])
    add_table_caption(doc,
        "Severity classification thresholds used by the rule-based Automated Interpretation Engine. "
        "All ML probability thresholds represent relative discriminative scores, not calibrated absolute risk estimates.")

    add_h2(doc, "Module 4 — Interactive ADMET Dashboard")

    add_body(doc,
        "The dashboard aggregates all tool outputs into a unified visual summary updated in real time as predictions "
        "resolve (Figure 3). Panels: (i) Summary metric cards (mean MW, LogP, QED, oral bioavailability, Lipinski "
        "compliance); (ii) Safety Flags — labelled progress bars for hERG, DILI, PAINS, BRENK, and BBB+ proportions "
        "across the batch; (iii) StopTox acute toxicity distribution; (iv) ESOL solubility distribution; "
        "(v) Per-Molecule Risk Matrix — colour-coded table with Overall/hERG/DILI/ClinTox/BBB/QED per molecule; "
        "(vi) CYP Inhibition Heatmap — 5-isoform × N-molecule probability matrix with three-tier colouring "
        "(green < 25%, amber 25–50%, red > 50%).")

    add_h2(doc, "Module 5 — Local Descriptor Engine and ESOL Solubility")

    add_body(doc,
        "Over 60 physicochemical and topological descriptors computed locally via RDKit: constitutional "
        "(MW, FractionCSP3, Labute ASA, MolMR); drug-likeness — QED [14], Lipinski [15], Ghose, Veber [12], "
        "Egan, Muegge; topological indices (Balaban J, BertzCT, Kappa 1–3, Chi series); electronic/VSA descriptors "
        "(PEOE_VSA, SMR_VSA, SlogP_VSA); and structural alerts via PAINS/BRENK/NIH catalogs.")

    add_body(doc, "Aqueous solubility is estimated via the ESOL QSAR model (Delaney, 2004 [16]):")
    add_esol_block(doc)

    add_body(doc,
        "ESOL provides four-category solubility classification (Soluble/Moderately/Poorly/Insoluble) with ±1 "
        "log-unit uncertainty. Four fingerprint protocols are exported in QSAR-ready format: RDKit (1,024 bits), "
        "Morgan/ECFP4 (2,048 bits, radius 2), MACCS keys (167 bits), Atom Pairs (2,048 bits).")

    add_h2(doc, "Module 6 — Batch Processing, Export, and Auxiliary Tools")

    add_body(doc,
        "CSV batch input (up to 500 compounds; Name + SMILES columns). Per-compound error isolation. Export: "
        "structured Excel workbook with ADMET comparison, flat records, and fingerprint matrices formatted for "
        "scikit-learn/DeepChem ingestion; PDF clinical summary. PepLink integration for bidirectional "
        "peptide-SMILES translation. Tanimoto similarity search (configurable Morgan radius 1–4; colour-coded "
        "Tc ≥ 0.70/0.40–0.70/< 0.40). SMILES-to-IUPAC via PubChem PUG REST API.")

    # ════════════════════════════════════════════════════════════════════════
    # RESULTS AND DISCUSSION
    # ════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Results and Discussion")
    add_h2(doc, "Benchmark Against Open-Source Cheminformatics Platforms")

    add_body(doc,
        "To characterise SMILESRender's position in the open-source cheminformatics landscape, we performed a "
        "systematic feature comparison against five representative platforms (Table 2). Platforms were evaluated "
        "based on publicly documented capabilities as of May 2026; features confirmed through direct testing are "
        "marked accordingly.")

    # ── TABLE 2 ───────────────────────────────────────────────────────────────
    p_cap_t2 = doc.add_paragraph()
    set_para_spacing(p_cap_t2, space_before=8, space_after=4)
    r_t2 = p_cap_t2.add_run(
        "Table 2. Systematic benchmark: SMILESRender vs. five open-source cheminformatics platforms.")
    r_t2.bold = True
    r_t2.font.name = 'Calibri'
    r_t2.font.size = Pt(10)

    t2_headers = ["Feature", "SMILESRender", "DataWarrior [1]", "KNIME + RDKit [2]",
                  "Galaxy Cheminf. [3]", "ChemMine [4]", "MarvinSketch (free)"]
    t2_rows = [
        ["Interface", "Web", "Desktop", "Desktop/Server", "Web (server)", "Web", "Desktop"],
        ["No installation required", "✓", "—", "—", "✓ᵃ", "✓", "—"],
        ["Docker/offline deployment", "✓", "—", "Partial", "—", "—", "—"],
        ["Open source (OSI license)", "✓ (MIT)", "Partialᵇ", "✓", "✓", "✓", "—"],
        ["2D structure rendering (batch)", "✓", "✓", "✓ᶜ", "—", "Partial", "✓"],
        ["Interactive structure editor", "✓ (JSME)", "✓", "—", "—", "—", "✓"],
        ["Reaction SMILES visualisation", "✓", "✓", "✓ᶜ", "—", "—", "✓"],
        ["ADMET prediction — local ML", "✓ (85+ endpoints)", "Partialᵈ", "—ᵉ", "—", "—", "Partialᶠ"],
        ["BBB permeability model", "✓", "—", "—", "—", "—", "—"],
        ["Tox21 multi-endpoint (12)", "✓", "—", "—", "—", "—", "—"],
        ["Chemprop D-MPNN (53 props)", "✓", "—", "—", "—", "—", "—"],
        ["External ADMET integration", "✓ (3 services)", "—", "Manualᵍ", "—", "—", "—"],
        ["Automated interpretation", "✓", "—", "—", "—", "—", "—"],
        ["Drug-likeness (Ro5/Veber/QED)", "✓", "✓", "✓ᶜ", "—", "Partial", "✓"],
        ["PAINS structural alerts", "✓", "—", "✓ᶜ", "—", "—", "—"],
        ["BRENK / NIH alerts", "✓", "—", "✓ᶜ", "—", "—", "—"],
        ["60+ RDKit descriptor panel", "✓", "✓ᵈ", "✓ᶜ", "Partial", "—", "Partial"],
        ["4 fingerprint types (ML-ready)", "✓", "Partial", "✓ᶜ", "—", "Partial", "—"],
        ["Chemical similarity search", "✓", "✓", "✓ᶜ", "—", "✓", "—"],
        ["IUPAC nomenclature", "✓", "Partial", "—", "—", "—", "✓"],
        ["Batch CSV upload", "✓", "✓", "✓", "Partial", "✓", "—"],
        ["CYP inhibition heatmap", "✓ (5 isoforms)", "—", "—", "—", "—", "—"],
        ["Per-molecule risk matrix", "✓", "—", "—", "—", "—", "—"],
        ["Required expertise", "Minimal", "Moderate", "High", "High", "Minimal", "Minimal"],
    ]
    build_styled_table(doc, t2_headers, t2_rows,
                       col_widths=[4.2, 2.5, 2.5, 2.8, 2.8, 2.2, 2.8])
    add_table_caption(doc,
        "ᵃ Galaxy requires account registration and server allocation. "
        "ᵇ DataWarrior is freely distributed but not open-source (source code not publicly available; "
        "Actelion proprietary licence). "
        "ᶜ Requires workflow construction in KNIME node editor. "
        "ᵈ DataWarrior computes physicochemical descriptors and basic drug-likeness but has no ML-based ADMET "
        "toxicity endpoints. "
        "ᵉ KNIME has no built-in ADMET prediction; external API nodes can be added but require individual service accounts. "
        "ᶠ MarvinSketch computes pKa, logP, and solubility via ChemAxon proprietary models; no toxicity endpoints. "
        "ᵍ KNIME can call external REST APIs but requires manual workflow configuration per service.")

    add_body_mixed(doc, [("Key findings from the benchmark:", True, False)])
    findings = [
        "SMILESRender is the only open-source platform providing integrated, web-native multi-engine ADMET "
        "profiling without programming expertise.",
        "DataWarrior is the closest general-purpose competitor for structure-based analysis but has no ADMET ML "
        "models, no external service integration, and no automated interpretation.",
        "KNIME is the most powerful but requires expert workflow configuration — it is a platform for building "
        "tools, not a ready-to-use tool itself.",
        "No competitor provides coverage of all five ADMET categories (A, D, M, E, T) in a single session.",
        "Only SMILESRender and MarvinSketch are accessible to non-computational users without programming or "
        "workflow configuration; MarvinSketch provides no ADMET prediction beyond logP/pKa/solubility.",
    ]
    for i, finding in enumerate(findings, 1):
        p = doc.add_paragraph(style='List Number')
        set_para_spacing(p, space_after=3)
        run = p.add_run(finding)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)

    add_h2(doc, "ADMET Coverage Comparison")

    add_body(doc,
        "Figure 4 presents a radar chart comparing ADMET endpoint coverage across the benchmarked platforms. "
        "SMILESRender covers all five ADMET categories with multiple endpoints per category; DataWarrior provides "
        "absorption-related physicochemical properties only; KNIME with RDKit nodes covers descriptors but no "
        "prediction endpoints; and the remaining platforms cover two or fewer categories.")

    add_body(doc,
        "The cross-tool consensus approach enabled by SMILESRender's architecture provides an additional analytical "
        "advantage: when multiple independent models (local GBM BBB, Chemprop BBB_Martins, ProTox 3.0 BBB) predict "
        "the same BBB class concordantly, the signal carries higher confidence than any single-model prediction. "
        "Similarly, when Tox21-RF flags hERG-associated stress response pathways (SR-ARE, SR-MMP) and Chemprop "
        "simultaneously predicts high hERG probability, the combination is a stronger early warning than either "
        "tool alone.")

    add_h2(doc, "Validation: Biological Plausibility Check with Ten FDA-Approved Drugs")

    add_body(doc,
        "To verify descriptor accuracy and biological plausibility of ADMET predictions — not to provide "
        "statistical model validation (n = 10 is insufficient for that purpose) — we evaluated ten marketed drugs "
        "spanning six therapeutic classes (Table 3).")

    # ── TABLE 3 ───────────────────────────────────────────────────────────────
    p_cap_t3 = doc.add_paragraph()
    set_para_spacing(p_cap_t3, space_before=8, space_after=4)
    r_t3 = p_cap_t3.add_run(
        "Table 3. Descriptor computation and ADMET plausibility check for ten FDA-approved drugs.")
    r_t3.bold = True
    r_t3.font.name = 'Calibri'
    r_t3.font.size = Pt(10)

    t3_headers = ["Drug", "Class", "MW", "LogP", "TPSA", "QED", "Ro5",
                  "ESOL", "BBB", "hERG %", "DILI %", "CYP flag", "Expected CNS"]
    t3_rows = [
        ["Aspirin",       "Analgesic",   "180.2", "1.19",  "63.6",  "0.55", "Pass",  "Soluble", "BBB+", "5",  "32", "—",      "Partial"],
        ["Ibuprofen",     "NSAID",       "206.3", "3.97",  "37.3",  "0.73", "Pass",  "Mod.",    "BBB+", "3",  "21", "—",      "Limited"],
        ["Acetaminophen", "Analgesic",   "151.2", "0.46",  "49.3",  "0.59", "Pass",  "Soluble", "BBB+", "2",  "18", "—",      "Yes"],
        ["Caffeine",      "CNS",         "194.2", "0.16",  "61.4",  "0.56", "Pass",  "Soluble", "BBB+", "5",  "38", "—",      "Yes ✓"],
        ["Metformin",     "Antidiabetic","129.2", "−1.43", "88.5",  "0.30", "Pass",  "Soluble", "BBB−", "1",  "12", "—",      "No ✓"],
        ["Atorvastatin",  "Statin",      "558.6", "5.67",  "111.8", "0.34", "Fail*", "Poorly",  "BBB−", "12", "45", "CYP3A4", "No ✓"],
        ["Sildenafil",    "PDE5-i",      "474.6", "2.77",  "113.0", "0.53", "Pass",  "Mod.",    "BBB−", "8",  "38", "CYP3A4", "No ✓"],
        ["Lisinopril",    "ACE-i",       "405.5", "−0.09", "138.9", "0.29", "Pass",  "Soluble", "BBB−", "3",  "22", "—",      "No ✓"],
        ["Tamoxifen",     "SERM",        "371.5", "6.30",  "41.6",  "0.44", "Fail‡", "Poorly",  "BBB+", "11", "55ʰ","CYP2D6", "Yes ✓"],
        ["Ciprofloxacin", "Antibiotic",  "331.3", "0.28",  "74.6",  "0.49", "Pass",  "Soluble", "BBB−", "4",  "29", "—",      "No ✓"],
    ]
    build_styled_table(doc, t3_headers, t3_rows,
                       col_widths=[2.5, 2.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.8, 1.5, 1.8, 1.8, 2.2, 2.5])
    add_table_caption(doc,
        "MW in g/mol; TPSA in Å²; ESOL: Soluble > −2, Mod. −4 to −2, Poorly < −4 mol/L; BBB: local GBM model; "
        "hERG/DILI: admet_ai Chemprop (%); CYP flag: isoform with highest inhibition probability ≥ 0.50. "
        "*MW > 500 g/mol. ‡LogP > 5. ʰ DILI 55% flagged high; Tamoxifen carries a black-box warning for "
        "hepatotoxicity in long-term use. ✓ indicates concordance with established clinical CNS profile.")

    add_body(doc,
        "BBB classification was concordant with established CNS profiles for 9/10 compounds. The interpretation "
        "engine correctly annotated Atorvastatin's Lipinski violation as \"MW violation (558.6 g/mol > 500); note: "
        "transported by OATP1B1/B3\" and Lisinopril's high TPSA as \"absorption risk (TPSA 138.9 Å²); note: PepT1 "
        "substrate — transporter-mediated absorption known\". The CYP3A4 flag for Atorvastatin and Sildenafil "
        "accurately reflects their documented metabolic routes. The Tox21-RF flagged Tamoxifen for NR-ER activity "
        "(estrogen receptor agonism), consistent with its mechanism of action.")

    add_h2(doc, "Batch Processing Case Study: Thieno[2,3-b]pyridine Library")

    add_body(doc,
        "A library of 20 thieno[2,3-b]pyridine derivatives (DADOS_Uminho_1) — a kinase inhibitor scaffold — was "
        "processed via batch CSV upload. Full ADMET profiling across all local models and external oracles completed "
        "in 11 min 45 s ± 1 min 12 s (three independent runs, single analyst experienced with all tools). The "
        "equivalent manual workflow — entering 20 SMILES into three external services, downloading and consolidating "
        "results — required 2 h 52 min ± 18 min.")

    add_body(doc,
        "Key outputs: 14/20 (70%) predicted BBB+ (all within training AD; Tanimoto NN 0.33–0.51); 3 PAINS alerts "
        "(rhodanine ×2, catechol ×1); 2 compounds with hERG probability > 0.60 (flagged high); CYP3A4 the most "
        "frequently inhibited isoform (11/20, 55%). The consolidated Excel export was immediately usable for SAR "
        "analysis without any data reformatting.")

    add_h2(doc, "Platform Accessibility and Reproducibility")

    add_body(doc,
        "SMILESRender requires no software installation for end users: the public cloud instance at "
        "https://smiles-render.onrender.com is accessible via any modern browser. For groups requiring data "
        "sovereignty or air-gapped environments, Docker Compose deployment takes under 5 minutes and produces a "
        "bit-identical local instance. The Docker image captures all library versions, model weights, and "
        "configuration — ensuring that a computational result described in a publication can be independently "
        "reproduced years later by simply running `docker compose up`.")

    # ════════════════════════════════════════════════════════════════════════
    # PLANNED EXTENSIONS
    # ════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Planned Extensions")

    add_body_mixed(doc, [
        ("3D Conformer Generation and Docking Interface:", True, False),
        (" A /generate/3d endpoint using RDKit ETKDG [17] and MMFF94 minimisation is in development, feeding "
         "into an AutoDock-GPU [18] / Vina [19] docking module with Meeko [20] receptor preparation and 3Dmol.js "
         "browser visualisation. This will be benchmarked on CASF-2016 re-docking tasks [21].", False, False)])

    add_body_mixed(doc, [
        ("Enhanced ADMET models:", True, False),
        (" Scaffold-disjoint (Bemis-Murcko) evaluation of the BBB model; consensus Tox21 model (RF + Chemprop "
         "multi-task); OPERA solubility integration [22] for higher-accuracy aqueous solubility predictions.", False, False)])

    # ════════════════════════════════════════════════════════════════════════
    # CONCLUSIONS
    # ════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Conclusions")

    add_body(doc,
        "SMILESRender provides what no existing open-source tool currently offers: a web-native, session-consistent, "
        "Docker-reproducible platform consolidating the full cheminformatics workflow — structure rendering, "
        "interactive editing, 85+ endpoint ADMET profiling, structural alert screening, descriptor computation, "
        "and automated plain-language interpretation — accessible without programming expertise.")

    add_body(doc,
        "A systematic benchmark against DataWarrior, KNIME, Galaxy, ChemMine, and MarvinSketch confirms that "
        "SMILESRender uniquely covers all five ADMET categories (Absorption, Distribution, Metabolism, Excretion, "
        "and Toxicity) through complementary local ML and external oracle layers, with automated cross-tool "
        "interpretation unavailable in any comparator platform. The embedded ADMET stack — covering hERG "
        "cardiotoxicity, DILI, 12 Tox21 bioassays, 5-isoform CYP metabolism, BBB permeability with applicability "
        "domain, and full organ-toxicity profiling — represents the most comprehensive ADMET coverage available "
        "in an open-source, offline-capable deployment.")

    add_body(doc,
        "The platform is freely available under the MIT license. All model weights and training scripts are "
        "distributed in the repository for full reproducibility.")

    # ════════════════════════════════════════════════════════════════════════
    # AVAILABILITY
    # ════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Availability and Requirements")

    avail_items = [
        ("Project name:", " SMILESRender"),
        ("Home page:", " https://github.com/rubithedev/smiles-render-web"),
        ("Cloud instance:", " https://smiles-render.onrender.com"),
        ("OS:", " Platform-independent; Docker recommended; tested on Linux Ubuntu 22.04 and Windows 11"),
        ("Languages:", " Python 3.12, TypeScript (React 19)"),
        ("Dependencies:", " Flask 3.0.3, RDKit 2024.3.6, scikit-learn 1.8, admet_ai ≥ 1.0, "
         "Waitress 3.0.1, Redis 7.4, Bun 1.1"),
        ("License:", " MIT"),
    ]
    for bold_part, rest in avail_items:
        p = doc.add_paragraph(style='List Bullet')
        set_para_spacing(p, space_after=3)
        add_run(p, bold_part, bold=True, size=10.5)
        add_run(p, rest, size=10.5)

    # ════════════════════════════════════════════════════════════════════════
    # ABBREVIATIONS
    # ════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Abbreviations")

    add_body(doc,
        "AD: applicability domain; ADMET: Absorption, Distribution, Metabolism, Excretion, Toxicity; "
        "AUC: area under the ROC curve; BBB: blood-brain barrier; BCS: Biopharmaceutics Classification System; "
        "CYP: cytochrome P450; D-MPNN: Directed Message Passing Neural Network; DDI: drug–drug interaction; "
        "DILI: drug-induced liver injury; ECFP: Extended Connectivity Fingerprint; ESOL: Estimated SOLubility; "
        "GBM: Gradient Boosting Machine; GHS: Globally Harmonized System; hERG: human Ether-à-go-go-Related Gene; "
        "HIA: human intestinal absorption; JSME: Java Structure Molecular Editor; ML: machine learning; "
        "MPO: multi-parameter optimisation; NCE: new chemical entity; PAINS: pan-assay interference compounds; "
        "PAMPA: parallel artificial membrane permeability assay; PPBR: plasma protein binding ratio; "
        "QSAR: quantitative structure–activity relationship; QED: quantitative estimate of drug-likeness; "
        "RF: Random Forest; RDKit: open-source cheminformatics toolkit; Ro5: Lipinski Rule of 5; "
        "SAR: structure–activity relationship; SMILES: Simplified Molecular Input Line Entry System; "
        "SOM: self-organising map; TDC: Therapeutics Data Commons; TPSA: topological polar surface area; "
        "VDss: volume of distribution at steady state.")

    # ════════════════════════════════════════════════════════════════════════
    # DECLARATIONS
    # ════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Declarations")

    p = doc.add_paragraph()
    set_para_spacing(p, space_after=4)
    add_run(p, "Competing interests: ", bold=True, size=10.5)
    add_run(p, "The authors declare no competing interests.", size=10.5)

    p = doc.add_paragraph()
    set_para_spacing(p, space_after=4)
    add_run(p, "Authors' contributions: ", bold=True, size=10.5)
    add_run(p, "RABS conceived and implemented the platform, trained all local ML models, performed all "
               "benchmarking, and drafted the manuscript. GG contributed to architecture design and manuscript "
               "revision. All authors approved the final manuscript.", size=10.5)

    p = doc.add_paragraph()
    set_para_spacing(p, space_after=4)
    add_run(p, "Acknowledgements: ", bold=True, size=10.5)
    add_run(p, "The authors thank the developers of RDKit (G. Landrum et al.), ", size=10.5)
    add_run(p, "admet_ai", italic=True, size=10.5)
    add_run(p, " (K. Swanson et al., Stanford), StopTox/StopLight (A. Borrel, N. Kleinstreuer et al., "
               "NIH/NIEHS), ProTox 3.0 (P. Banerjee et al., Charité), DataWarrior (T. Sander, "
               "Actelion/Sanofi), B3DB (F. Meng et al.), and the Tox21 Challenge (NIH). The "
               "thieno[2,3-b]pyridine dataset (DADOS_Uminho_1) was used with permission.", size=10.5)

    # ════════════════════════════════════════════════════════════════════════
    # REFERENCES
    # ════════════════════════════════════════════════════════════════════════
    add_h1(doc, "References")

    refs = [
        # (num, [(text, bold, italic), ...])
        (1, [("1. ", False, False),
             ("Sander T, Freyss J, von Korff M, Rufener C. DataWarrior: an open-source program for chemistry "
              "aware data visualization and analysis. ", False, False),
             ("J Chem Inf Model.", False, True),
             (" 2015;55(2):460–473. https://doi.org/10.1021/ci500588j", False, False)]),
        (2, [("2. ", False, False),
             ("Berthold MR, Cebron N, Dill F, Gabriel TR, Kötter T, Meinl T, et al. KNIME — the Konstanz "
              "information miner: version 2.0 and beyond. ", False, False),
             ("ACM SIGKDD Explor Newsl.", False, True),
             (" 2009;11(1):26–31. https://doi.org/10.1145/1656274.1656280", False, False)]),
        (3, [("3. ", False, False),
             ("Afgan E, Baker D, Batut B, van den Beek M, Bouvier D, Čech M, et al. The Galaxy platform for "
              "accessible, reproducible and collaborative biomedical analyses: 2018 update. ", False, False),
             ("Nucleic Acids Res.", False, True),
             (" 2018;46(W1):W537–W544. https://doi.org/10.1093/nar/gky379", False, False)]),
        (4, [("4. ", False, False),
             ("Backman TWH, Cao Y, Girke T. ChemMine Tools: an online service for analyzing and clustering "
              "small molecules. ", False, False),
             ("Nucleic Acids Res.", False, True),
             (" 2011;39(Web Server issue):W486–W491. https://doi.org/10.1093/nar/gkr492", False, False)]),
        (5, [("5. ", False, False),
             ("Landrum G, Tosco P, Kelley B, et al. RDKit: open-source cheminformatics. Version 2024.03.6. "
              "https://doi.org/10.5281/zenodo.591637", False, False)]),
        (6, [("6. ", False, False),
             ("Ertl P, Bienfait B. JSME: a free molecule editor in JavaScript. ", False, False),
             ("J Cheminform.", False, True),
             (" 2013;5:24. https://doi.org/10.1186/1758-2946-5-24", False, False)]),
        (7, [("7. ", False, False),
             ("Yang K, Swanson K, Jin W, Coley C, Eiden P, Gao H, et al. Analyzing learned molecular "
              "representations for property prediction. ", False, False),
             ("J Chem Inf Model.", False, True),
             (" 2019;59(8):3370–3388. https://doi.org/10.1021/acs.jcim.9b00237", False, False)]),
        (8, [("8. ", False, False),
             ("Swanson K, Boros P, Chen LC, Bhatt DL, Bonn-Miller MO, Wang H, Plotkin SS. ADMET-AI: a machine "
              "learning ADMET platform for evaluation of large-scale chemical libraries. ", False, False),
             ("Bioinformatics.", False, True),
             (" 2024;40(7):btae416. https://doi.org/10.1093/bioinformatics/btae416", False, False)]),
        (9, [("9. ", False, False),
             ("Borrel A, Mansouri K, Nolte S, Zurlinden T, Huang R, Xia M, Houck KA, Kleinstreuer NC. StopTox: "
              "an in silico alternative to animal acute systemic toxicity tests. ", False, False),
             ("Environ Health Perspect.", False, True),
             (" 2022;130(2):027014. https://doi.org/10.1289/EHP9341", False, False)]),
        (10, [("10. ", False, False),
              ("Borrel A, Huang R, Sakamuru S, Xia M, Simeonov A, Mansouri K, Kleinstreuer NC. "
               "High-throughput screening to predict chemical-assay interference. ", False, False),
              ("Sci Rep.", False, True),
              (" 2020;10(1):3986. https://doi.org/10.1038/s41598-020-60747-3", False, False)]),
        (11, [("11. ", False, False),
              ("Banerjee P, Dehnbostel FO, Preissner R. ProTox-3.0: a webserver for the prediction of toxicity "
               "of chemicals. ", False, False),
              ("Nucleic Acids Res.", False, True),
              (" 2024;52(W1):W513–W520. https://doi.org/10.1093/nar/gkae303", False, False)]),
        (12, [("12. ", False, False),
              ("Veber DF, Johnson SR, Cheng HY, Smith BR, Ward KW, Kopple KD. Molecular properties that "
               "influence the oral bioavailability of drug candidates. ", False, False),
              ("J Med Chem.", False, True),
              (" 2002;45(12):2615–2623. https://doi.org/10.1021/jm020017n", False, False)]),
        (13, [("13. ", False, False),
              ("Baell JB, Holloway GA. New substructure filters for removal of pan assay interference compounds "
               "(PAINS) from screening libraries and for their exclusion in bioassays. ", False, False),
              ("J Med Chem.", False, True),
              (" 2010;53(7):2719–2740. https://doi.org/10.1021/jm901137j", False, False)]),
        (14, [("14. ", False, False),
              ("Bickerton GR, Paolini GV, Besnard J, Muresan S, Hopkins AL. Quantifying the chemical beauty "
               "of drugs. ", False, False),
              ("Nat Chem.", False, True),
              (" 2012;4(2):90–98. https://doi.org/10.1038/nchem.1243", False, False)]),
        (15, [("15. ", False, False),
              ("Lipinski CA, Lombardo F, Dominy BW, Feeney PJ. Experimental and computational approaches to "
               "estimate solubility and permeability in drug discovery and development settings. ", False, False),
              ("Adv Drug Deliv Rev.", False, True),
              (" 2001;46(1–3):3–26. https://doi.org/10.1016/s0169-409x(00)00129-0", False, False)]),
        (16, [("16. ", False, False),
              ("Delaney JS. ESOL: estimating aqueous solubility directly from molecular structure. ", False, False),
              ("J Chem Inf Comput Sci.", False, True),
              (" 2004;44(3):1000–1005. https://doi.org/10.1021/ci034243x", False, False)]),
        (17, [("17. ", False, False),
              ("Riniker S, Landrum GA. Better informed distance geometry: using what we know to improve "
               "conformation generation. ", False, False),
              ("J Chem Inf Model.", False, True),
              (" 2015;55(12):2562–2574. https://doi.org/10.1021/acs.jcim.5b00654", False, False)]),
        (18, [("18. ", False, False),
              ("Santos-Martins D, Solis-Vasquez L, Tillack AF, Sanner MF, Koch A, Forli S. Accelerating "
               "AutoDock4 with GPUs and gradient-based local search. ", False, False),
              ("J Chem Theory Comput.", False, True),
              (" 2021;17(2):1060–1073. https://doi.org/10.1021/acs.jctc.0c01006", False, False)]),
        (19, [("19. ", False, False),
              ("Eberhardt J, Santos-Martins D, Tillack AF, Forli S. AutoDock Vina 1.2.0: new docking methods, "
               "expanded force field, and Python bindings. ", False, False),
              ("J Chem Inf Model.", False, True),
              (" 2021;61(8):3891–3898. https://doi.org/10.1021/acs.jcim.1c00203", False, False)]),
        (20, [("20. ", False, False),
              ("Forli S, Huey R, Pique ME, Sanner MF, Goodsell DS, Olson AJ. Computational protein–ligand "
               "docking and virtual drug screening with the AutoDock suite. ", False, False),
              ("Nat Protoc.", False, True),
              (" 2016;11(5):905–919. https://doi.org/10.1038/nprot.2016.051", False, False)]),
        (21, [("21. ", False, False),
              ("Su M, Yang Q, Du Y, Feng G, Liu Z, Li Y, Wang R. Comparative assessment of scoring functions: "
               "the CASF-2016 and CASF-2013 benchmarks. ", False, False),
              ("J Chem Inf Model.", False, True),
              (" 2019;59(2):895–913. https://doi.org/10.1021/acs.jcim.8b00545", False, False)]),
        (22, [("22. ", False, False),
              ("Mansouri K, Grulke CM, Judson RS, Williams AJ. OPERA models for predicting physicochemical "
               "properties and environmental fate endpoints. ", False, False),
              ("J Cheminform.", False, True),
              (" 2018;10(1):10. https://doi.org/10.1186/s13321-018-0263-1", False, False)]),
        (23, [("23. ", False, False),
              ("Meng F, Xi Y, Huang J, Ayers PW. A curated diverse molecular database of blood-brain barrier "
               "permeability with chemical descriptors. ", False, False),
              ("Sci Data.", False, True),
              (" 2021;8(1):289. https://doi.org/10.1038/s41597-021-01069-5", False, False)]),
        (24, [("24. ", False, False),
              ("Tice RR, Austin CP, Kavlock RJ, Bucher JR. Improving the human hazard characterization of "
               "chemicals: a Tox21 update. ", False, False),
              ("Environ Health Perspect.", False, True),
              (" 2013;121(7):756–765. https://doi.org/10.1289/ehp.1205784", False, False)]),
        (25, [("25. ", False, False),
              ("Dhanjal JK, Wang S, Bhinder B, Singh Y, Kaur H, Grover A. GraphB3: an explainable graph "
               "convolutional network approach for blood-brain barrier permeability prediction. ", False, False),
              ("J Cheminform.", False, True),
              (" 2024;16:34. https://doi.org/10.1186/s13321-024-00831-4", False, False)]),
        (26, [("26. ", False, False),
              ("Muratov EN, Bajorath J, Sheridan RP, Tetko IV, Filimonov D, Poroikov V, et al. QSAR without "
               "borders. ", False, False),
              ("Chem Soc Rev.", False, True),
              (" 2020;49(11):3525–3564. https://doi.org/10.1039/d0cs00098a", False, False)]),
        (27, [("27. ", False, False),
              ("Waring MJ, Arrowsmith J, Leach AR, Leeson PD, Mandrell S, Owen RM, et al. An analysis of the "
               "attrition of drug candidates from four major pharmaceutical companies. ", False, False),
              ("Nat Rev Drug Discov.", False, True),
              (" 2015;14(7):475–486. https://doi.org/10.1038/nrd4609", False, False)]),
        (28, [("28. ", False, False),
              ("Daina A, Michielin O, Zoete V. SwissADME: a free web tool to evaluate pharmacokinetics, "
               "drug-likeness and medicinal chemistry friendliness of small molecules. ", False, False),
              ("Sci Rep.", False, True),
              (" 2017;7:42717. https://doi.org/10.1038/srep42717", False, False)]),
        (29, [("29. ", False, False),
              ("Gui C, Luo M, Wang Z, Ma H, Du Z, Yao L, et al. ADMETlab 3.0: an updated comprehensive online "
               "ADMET prediction tool with improved models and functions. ", False, False),
              ("Nucleic Acids Res.", False, True),
              (" 2024;52(W1):W197–W204. https://doi.org/10.1093/nar/gkae420", False, False)]),
        (30, [("30. ", False, False),
              ("Paul SM, Mytelka DS, Dunwiddie CT, Persinger CC, Munos BH, Lindborg SR, Schacht AL. How to "
               "improve R&D productivity: the pharmaceutical industry's grand challenge. ", False, False),
              ("Nat Rev Drug Discov.", False, True),
              (" 2010;9(3):203–214. https://doi.org/10.1038/nrd3078", False, False)]),
    ]

    for num, parts in refs:
        add_reference(doc, num, parts)

    # ════════════════════════════════════════════════════════════════════════
    # FIGURE LEGENDS
    # ════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Figure Legends")

    add_figure_legend(doc, 1,
        "SMILESRender hybrid system architecture. Left (teal): local in-process computation layer — three "
        "embedded ML models (Tox21-RF, BBB-GBM with AD flag, DeepADMET/Chemprop) running without network "
        "dependency. Right (amber): external oracle orchestration layer — StopTox, StopLight, ProTox 3.0, "
        "each fault-isolated in a ToolErrorBoundary. Top: React 19 frontend. Centre: Redis cache, Celery "
        "worker, Flask REST API. The local layer guarantees a minimum viable 85-endpoint ADMET profile at "
        "100% availability; the external layer adds supplementary coverage when network is available.")

    add_figure_legend(doc, 2,
        "SMILESRender module overview. Six first-class analytical modules accessible from the hub: "
        "(A) ADMET Profiling — multi-engine orchestration with real-time dashboard; "
        "(B) Molecular Renderer — batch PNG/ZIP export with transparent background; "
        "(C) Descriptors — 60+ local RDKit descriptors with fingerprint export; "
        "(D) Similarity Search — Tanimoto-ranked molecular comparison; "
        "(E) IUPAC Converter — PubChem-backed nomenclature; "
        "(F) Peptide Engineering — bidirectional peptide-SMILES translation.")

    add_figure_legend(doc, 3,
        "Interactive ADMET Dashboard for a representative 5-compound batch. Top row: summary metric cards. "
        "Middle-left: Safety Flags panel (hERG, DILI, PAINS, BRENK, BBB+ proportions as progress bars). "
        "Middle-right: StopTox acute toxicity distribution and ESOL solubility stacked bar. "
        "Bottom-left: Per-Molecule Risk Matrix (colour-coded Overall/hERG/DILI/ClinTox/BBB/QED per molecule). "
        "Bottom-right: CYP Inhibition Heatmap (5 isoforms × N compounds; green < 25%, amber 25–50%, red > 50%).")

    add_figure_legend(doc, 4,
        "ADMET endpoint coverage radar chart comparing SMILESRender with benchmarked open-source platforms. "
        "Five axes: Absorption, Distribution, Metabolism, Excretion, Toxicity. Each axis scored 0–100% based "
        "on number of endpoints covered relative to the maximum available across all platforms. SMILESRender "
        "achieves ≥ 80% on all five axes; DataWarrior covers Absorption only (physicochemical); KNIME+RDKit "
        "covers Absorption and limited Toxicity (structural alerts); ChemMine and Galaxy cover no ADMET axes.")

    add_figure_legend(doc, 5,
        "Batch processing workflow for 20 thieno[2,3-b]pyridine derivatives. "
        "(A) CSV upload and SMILES input interface with JSME structure editor. "
        "(B) Real-time prediction progress — tool status indicators for 6 engines (3 local, 3 external). "
        "(C) Consolidated ADMET dashboard on batch completion. "
        "(D) Excel export with per-compound ADMET comparison, flat records, and fingerprint matrix sheets.")

    # ── Save ──────────────────────────────────────────────────────────────────
    output_path = r"C:\Users\ruiab\Documents\SmileRender\SMILESRender_JCheminform_2026_v5.docx"
    doc.save(output_path)
    print(f"Saved: {output_path}")
    print(f"Size:  {os.path.getsize(output_path):,} bytes  "
          f"({os.path.getsize(output_path)/1024:.1f} KB)")
    return output_path


if __name__ == "__main__":
    build_document()
